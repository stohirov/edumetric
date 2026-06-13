#!/usr/bin/env python3
"""Post-process the built .docx:
  - build a BTEC-template-style title page (centred, sized) on its own page
  - visible borders + shaded, bold, repeating header rows on content tables
  - page break before each front-matter section, chapter (H1) and Appendix
  - centred figures and figure captions
"""
import pathlib
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = pathlib.Path(__file__).resolve().parent.parent / "EduMetric_Dissertation.docx"

BORDER_COLOR = "8C8C8C"
TITLE_BLUE = RGBColor(0x1F, 0x38, 0x64)   # matches Heading 1 colour
HEADER_FILL = "D9E1EC"   # light blue-grey, matches the chart palette


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")        # 0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)
        borders.append(el)
    # drop any existing tblBorders then add ours
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(borders)


def mark_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def first_row_is_header(table):
    cells = table.rows[0].cells
    return any(c.text.strip() for c in cells)


def style_tables(doc):
    n = 0
    for t in doc.tables:
        set_table_borders(t)
        # light cell padding
        if first_row_is_header(t):
            hdr = t.rows[0]
            mark_header_repeat(hdr)
            for c in hdr.cells:
                set_cell_shading(c, HEADER_FILL)
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
        n += 1
    return n


def _title_para(doc, text, size, *, bold=False, italic=False, color=None,
                align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4):
    """Create a title-page paragraph (appended; relocated by build_title_page)."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p._p


def _no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(borders)


def build_title_page(doc):
    """Rebuild the BTEC-template title page from scratch and place it first."""
    els = []
    els.append(_title_para(doc, "", 11, after=20))                       # top gap
    els.append(_title_para(doc, "PDP UNIVERSITY", 24, bold=True))
    els.append(_title_para(doc, "Faculty of Business Information Technology", 13))
    els.append(_title_para(doc, "Tashkent, Uzbekistan", 11, after=16))
    els.append(_title_para(doc, "INDEPENDENT PROJECT", 18, bold=True))
    els.append(_title_para(doc, "Pearson BTEC Level 6 Diploma in Digital Technologies", 11))
    els.append(_title_para(doc, "Unit 2 — Unit Code: 70726U — Credit Value: 30", 11, after=30))
    els.append(_title_para(doc, "EduMetric", 32, bold=True, color=TITLE_BLUE))
    els.append(_title_para(
        doc, "A Web-Based Multi-Dimensional Student Performance Analytics System",
        15, bold=True))
    els.append(_title_para(
        doc, "Design, Development and Evaluation of a Transparent "
        "Composite-Score Platform for Educational Institutions",
        12, italic=True, after=30))

    # student-info block — borderless 2-column table, bold labels
    fields = [
        ("Student Name:", "⟨STUDENT INPUT: Full Name⟩"),
        ("Student ID:", "⟨STUDENT INPUT: ID Number⟩"),
        ("Programme / Group:", "⟨STUDENT INPUT: BIT — Group Name⟩"),
        ("Project Format:", "Thesis-style ☐    Capstone-style ☒"),
        ("Supervisor:", "⟨STUDENT INPUT: Supervisor Full Name⟩"),
        ("Submission Date:", "⟨STUDENT INPUT: Day Month Year⟩"),
        ("Word Count:", "⟨STUDENT INPUT: XXXX words⟩"),
    ]
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _no_table_borders(tbl)
    for (label, value), row in zip(fields, tbl.rows):
        lc, vc = row.cells
        lc.width = Inches(2.3)
        vc.width = Inches(3.7)
        lp, vp = lc.paragraphs[0], vc.paragraphs[0]
        lr = lp.add_run(label); lr.font.bold = True; lr.font.size = Pt(11)
        vr = vp.add_run(value); vr.font.size = Pt(11)
    els.append(tbl._tbl)

    els.append(_title_para(doc, "", 11, after=8))                        # gap
    els.append(_title_para(
        doc, "Submitted in partial fulfilment of the requirements for the", 10))
    els.append(_title_para(
        doc, "Bachelor's Degree in Business Information Technology", 10, bold=True))
    els.append(_title_para(
        doc, "Academic Year ⟨STUDENT INPUT: 20XX–20XX⟩", 10))

    # relocate the freshly-appended elements to the very top of the body
    body = doc.element.body
    anchor = body[0]
    for el in els:
        body.remove(el)
    for el in els:
        anchor.addprevious(el)


def has_drawing(paragraph):
    return paragraph._p.findall(qn("w:r") + "/" + qn("w:drawing")) or \
        ".//" and paragraph._p.xpath(".//w:drawing")


def add_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pPr.append(pb)


def style_layout(doc):
    figs, breaks = 0, 0
    seen_chapter = False   # front-matter (H2) runs before the first H1 chapter
    for p in doc.paragraphs:
        style = p.style.name
        txt = p.text.strip()
        # page breaks: every chapter (H1), every front-matter section (H2
        # before the first chapter) and every Appendix start on a fresh page.
        if style == "Heading 1":
            seen_chapter = True
            add_page_break_before(p); breaks += 1
        elif style == "Heading 2" and (not seen_chapter or
                                       txt.startswith("Appendix ")):
            add_page_break_before(p); breaks += 1
        # centre figure images
        if p._p.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figs += 1
        # centre + shrink figure captions
        elif txt.startswith("Figure ") and "—" in txt and len(txt) < 220:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return figs, breaks


def main():
    doc = docx.Document(str(DOCX))
    nt = style_tables(doc)            # style content tables (before title table exists)
    build_title_page(doc)            # insert BTEC-style title page at the top
    nf, nb = style_layout(doc)
    doc.save(str(DOCX))
    print(f"styled {nt} tables, built title page, centred {nf} figures, "
          f"{nb} page breaks added")


if __name__ == "__main__":
    main()
